"""
Resume processing utilities for text extraction from various formats
"""
import os
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document
from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    try:
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        extracted = '\n'.join(text)
        logger.info("pdf_text_extracted", length=len(extracted), pages=len(text))
        return extracted.strip()
        
    except Exception as e:
        logger.error("pdf_extraction_failed", error=str(e), file=file_path)
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        doc = Document(file_path)
        text = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        extracted = '\n'.join(text)
        logger.info("docx_text_extracted", length=len(extracted), paragraphs=len(doc.paragraphs))
        return extracted.strip()
        
    except Exception as e:
        logger.error("docx_extraction_failed", error=str(e), file=file_path)
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a plain text file
    
    Args:
        file_path: Path to the text file
        
    Returns:
        File content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        logger.info("txt_text_extracted", length=len(text))
        return text.strip()
        
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            logger.info("txt_text_extracted_latin1", length=len(text))
            return text.strip()
        except Exception as e:
            logger.error("txt_extraction_failed", error=str(e), file=file_path)
            raise ValueError(f"Failed to read text file: {str(e)}")
    except Exception as e:
        logger.error("txt_extraction_failed", error=str(e), file=file_path)
        raise ValueError(f"Failed to read text file: {str(e)}")


def extract_resume_text(file_path: str, filename: str) -> str:
    """
    Extract text from resume file based on extension
    
    Args:
        file_path: Path to the resume file
        filename: Original filename with extension
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is not supported
    """
    file_extension = Path(filename).suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def validate_resume_file(filename: str, file_size: int, allowed_formats: list, max_size_bytes: int) -> tuple[bool, Optional[str]]:
    """
    Validate resume file before processing
    
    Args:
        filename: Original filename
        file_size: File size in bytes
        allowed_formats: List of allowed extensions (without dot)
        max_size_bytes: Maximum allowed file size
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    # Check file extension
    file_extension = Path(filename).suffix.lower().lstrip('.')
    if file_extension not in allowed_formats:
        return False, f"File format not allowed. Allowed formats: {', '.join(allowed_formats)}"
    
    # Check file size
    if file_size > max_size_bytes:
        max_size_mb = max_size_bytes / (1024 * 1024)
        return False, f"File size exceeds maximum allowed size of {max_size_mb}MB"
    
    return True, None


def sanitize_filename(filename: str) -> str:
    """
    Sanitize filename to prevent directory traversal and other issues
    
    Args:
        filename: Original filename
        
    Returns:
        Sanitized filename
    """
    # Get just the filename without path
    filename = os.path.basename(filename)
    
    # Remove or replace problematic characters
    invalid_chars = '<>:"|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    
    # Limit length
    if len(filename) > 255:
        name, ext = os.path.splitext(filename)
        filename = name[:250] + ext
    
    return filename
